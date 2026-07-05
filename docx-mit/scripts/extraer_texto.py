#!/usr/bin/env python3
"""
Extrae el texto de un .docx (párrafos y tablas) con python-docx. Reemplaza al binario
`extract-text` del sandbox de Anthropic, que no existe fuera de él. Implementación propia (MIT).

Uso:
    python -X utf8 extraer_texto.py documento.docx

Correr SIEMPRE con `-X utf8` en Windows.
"""
import sys

from docx import Document


def main():
    if len(sys.argv) != 2:
        print("uso: python -X utf8 extraer_texto.py <documento.docx>")
        sys.exit(1)

    doc = Document(sys.argv[1])
    for p in doc.paragraphs:
        print(p.text)
    for i, tabla in enumerate(doc.tables, 1):
        print(f"\n[Tabla {i}]")
        for fila in tabla.rows:
            print("\t".join(c.text for c in fila.cells))


if __name__ == "__main__":
    main()
