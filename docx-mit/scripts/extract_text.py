#!/usr/bin/env python3
"""
Extracts the text of a .docx (paragraphs and tables) with python-docx. Replaces the
sandbox-only `extract-text` binary. Own implementation (MIT).

Usage:
    python -X utf8 extract_text.py document.docx

Always run with `-X utf8` on Windows.
"""
import sys

from docx import Document


def main():
    if len(sys.argv) != 2:
        print("usage: python -X utf8 extract_text.py <document.docx>")
        sys.exit(1)

    doc = Document(sys.argv[1])
    for p in doc.paragraphs:
        print(p.text)
    for i, table in enumerate(doc.tables, 1):
        print(f"\n[Table {i}]")
        for row in table.rows:
            print("\t".join(c.text for c in row.cells))


if __name__ == "__main__":
    main()
