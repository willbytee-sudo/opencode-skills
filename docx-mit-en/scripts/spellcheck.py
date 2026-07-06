#!/usr/bin/env python3
"""
Proofreads the text of a .docx and FLAGS likely mistakes so the model sees them and fixes
them before delivering. Extracts the text with python-docx and applies light heuristics:
encoding artifacts (mojibake), doubled words, and space-before-punctuation. Then it prints
the full text for a visual proofread. Own implementation (MIT).

Usage:
    python -X utf8 spellcheck.py document.docx

Always run with `-X utf8` on Windows (otherwise the cp1252 console corrupts non-ASCII text).

This is an ASSISTANT, not an infallible checker: it surfaces suspects and shows the text;
the final call (and reading the rest for errors) is the model's.
"""
import re
import sys

# Encoding artifacts (UTF-8 decoded wrong) that often appear in broken text.
MOJIBAKE = ["Ã", "Â", "Ð", "�", "â€", "Ã©", "Ã±", "Ã¡", "Ã³", "Ãº"]
# Doubled consecutive words: "the the", "is is".
DOUBLED = re.compile(r"\b(\w+)\s+\1\b", re.IGNORECASE)
# A space before ,.;:!? (typographic slip).
SPACE_BEFORE_PUNCT = re.compile(r"\s+[,.;:!?]")


def check(text: str):
    """Return a list of (line_number, finding)."""
    findings = []
    for n, line in enumerate(text.splitlines(), 1):
        for mark in MOJIBAKE:
            if mark in line:
                findings.append((n, f"mojibake '{mark}' (corrupted text): {line.strip()[:80]}"))
                break
        for m in DOUBLED.finditer(line):
            # allow legitimate repeats like "had had", "that that" only rarely; still flag for review
            findings.append((n, f"doubled word '{m.group(0)}': {line.strip()[:80]}"))
        if SPACE_BEFORE_PUNCT.search(line):
            findings.append((n, f"space before punctuation: {line.strip()[:80]}"))
    # de-duplicate keeping order
    seen, unique = set(), []
    for f in findings:
        if f not in seen:
            seen.add(f); unique.append(f)
    return unique


def extract_text_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) != 2:
        print("usage: python -X utf8 spellcheck.py <document.docx>")
        sys.exit(1)

    text = extract_text_docx(sys.argv[1])
    findings = check(text)

    print("=" * 70)
    print(f"PROOFREAD — {len(findings)} possible issue(s) flagged")
    print("=" * 70)
    if findings:
        for n, msg in findings:
            print(f"  line {n}: {msg}")
    else:
        print("  (no automatic suspects — still READ the text below)")

    print("\n" + "-" * 70)
    print("FULL TEXT (read it for mistakes the heuristics can't catch):")
    print("-" * 70)
    print(text)


if __name__ == "__main__":
    main()
